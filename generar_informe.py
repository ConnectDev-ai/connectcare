# -*- coding: utf-8 -*-
"""
generar_informe.py — Genera Geoworkshop_Arquitectura.docx
Ejecutar desde la raíz del proyecto:
    python generar_informe.py
"""
import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────
# COLORES
# ─────────────────────────────────────────────
C_DARK   = (0.11, 0.11, 0.11)
C_CARD   = (0.14, 0.13, 0.13)
C_GOLD   = (0.949, 0.745, 0.549)
C_BLUE   = (0.655, 0.800, 0.918)
C_GREEN  = (0.525, 0.937, 0.671)
C_PURPLE = (0.769, 0.710, 0.992)
C_RED    = (0.973, 0.506, 0.506)
C_TEAL   = (0.431, 0.906, 0.718)

# ─────────────────────────────────────────────
# HELPER — celda con fondo de color
# ─────────────────────────────────────────────
def _cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[i])

def _heading_row(table, *texts, bg="1C1B1B", fg="F2BE8C"):
    row = table.rows[0]
    for i, txt in enumerate(texts):
        cell = row.cells[i]
        _cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(fg)

# ─────────────────────────────────────────────
# DIAGRAMA DE FLUJO
# ─────────────────────────────────────────────
def _draw_flow() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_facecolor(C_DARK)
    fig.patch.set_facecolor(C_DARK)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")

    def box(x, y, w, h, label, sublabel="", color=C_CARD, text_color=C_GOLD, fontsize=9):
        rect = mpatches.FancyBboxPatch(
            (x, y), w, h,
            boxstyle="round,pad=0.05",
            facecolor=color, edgecolor=(0.3, 0.27, 0.27), linewidth=0.8,
        )
        ax.add_patch(rect)
        cy = y + h / 2 + (0.1 if sublabel else 0)
        ax.text(x + w / 2, cy, label,
                ha="center", va="center",
                color=text_color, fontsize=fontsize, fontweight="bold",
                wrap=True, multialignment="center")
        if sublabel:
            ax.text(x + w / 2, y + h / 2 - 0.22, sublabel,
                    ha="center", va="center",
                    color=(0.7, 0.65, 0.6), fontsize=6.5)

    def arrow(x1, y1, x2, y2, color=(0.5, 0.45, 0.4)):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=1.2))

    def label_arrow(x, y, txt):
        ax.text(x, y, txt, color=(0.55, 0.5, 0.45), fontsize=6.5, ha="center")

    # ── FILA 1: Fuentes de datos ──────────────────────────────────────────
    ax.text(7, 8.65, "FUENTES DE DATOS EXTERNAS", ha="center", color=C_GOLD,
            fontsize=8, fontweight="bold", alpha=0.6)

    box(0.3, 7.6, 2.6, 0.85, "Copiloto API", "GPS / CSV vehicular", color=(0.10, 0.16, 0.22), text_color=C_BLUE)
    box(3.3, 7.6, 2.6, 0.85, "Geotab API", "3 bases: CO / CL / buses", color=(0.10, 0.16, 0.22), text_color=C_BLUE)
    box(6.3, 7.6, 2.6, 0.85, "SAP ERP Kaufmann", "smart_contract_vinSet", color=(0.10, 0.16, 0.22), text_color=C_BLUE)
    box(9.3, 7.6, 2.6, 0.85, "NHTSA / WMI", "Decode VIN", color=(0.10, 0.16, 0.22), text_color=C_BLUE)
    box(12.2, 7.6, 1.5, 0.85, "Excel\nmaster_Flota\ntalleres", color=(0.14, 0.18, 0.14), text_color=C_GREEN, fontsize=7)

    # ── FILA 2: Pipeline ─────────────────────────────────────────────────
    ax.text(4.5, 7.05, "PIPELINE DIARIO", ha="center", color=C_GOLD,
            fontsize=8, fontweight="bold", alpha=0.6)
    box(0.3, 5.85, 5.7, 1.0, "app.py  —  Pipeline principal",
        "GitHub Actions · cron 14:00 UTC · Ubuntu · Python 3.11",
        color=(0.17, 0.14, 0.10), text_color=C_GOLD, fontsize=9)

    # sub-pasos dentro del pipeline
    sub_y = 5.95
    sub_h = 0.5
    sub_w = 1.1
    sub_steps = [
        ("Auth\nCopiloto", (0.32, sub_y)),
        ("Fetch GPS\nCSV", (1.52, sub_y)),
        ("Geotab\nMultiDB", (2.72, sub_y)),
        ("Enrich\nMaster/SAP", (3.92, sub_y)),
        ("Haversine\n100 km", (5.12, 5.95)),
    ]
    for lbl, (sx, sy) in sub_steps:
        box(sx, sy, sub_w - 0.05, sub_h, lbl,
            color=(0.22, 0.18, 0.12), text_color=(0.9, 0.82, 0.7), fontsize=6.5)

    # build_sap_cache
    box(6.5, 5.85, 2.5, 1.0, "build_sap_cache.py",
        "Script manual · actualiza sap_vin_cache.json",
        color=(0.15, 0.12, 0.18), text_color=C_PURPLE, fontsize=8)

    # GitHub Actions badge
    box(9.3, 5.85, 4.4, 1.0, "GitHub Actions",
        "Secrets: COPILOTO, GEOTAB, PGPASSWORD, ERP_KEY",
        color=(0.12, 0.15, 0.12), text_color=C_GREEN, fontsize=8)

    # ── FILA 3: Base de datos ─────────────────────────────────────────────
    ax.text(3.5, 5.5, "ALMACENAMIENTO", ha="center", color=C_GOLD,
            fontsize=8, fontweight="bold", alpha=0.6)
    box(0.3, 4.2, 7.2, 1.1, "PostgreSQL / Neon  (PostGIS habilitado)",
        "snapshot_run · snapshot_unit · snapshot_taller_overlap · snapshot_taller_exclusive\n"
        "dim_taller · maintenance_ticket · maintenance_ticket_note · maintenance_record",
        color=(0.10, 0.12, 0.18), text_color=C_BLUE, fontsize=8)
    box(7.8, 4.2, 1.8, 1.1, "sap_vin_cache\n.json", color=(0.15, 0.12, 0.18), text_color=C_PURPLE, fontsize=7.5)
    box(9.9, 4.2, 1.7, 1.1, "vin_cache\n.json\n(NHTSA)", color=(0.10, 0.12, 0.18), text_color=C_BLUE, fontsize=7.5)
    box(11.8, 4.2, 1.9, 1.1, "Scripts/out/\n*.csv\n(snapshots locales)", color=(0.12, 0.15, 0.12), text_color=C_GREEN, fontsize=7)

    # ── FILA 4: API / Backend ─────────────────────────────────────────────
    ax.text(4.5, 3.85, "CAPA API / BACKEND", ha="center", color=C_GOLD,
            fontsize=8, fontweight="bold", alpha=0.6)
    box(0.3, 2.7, 5.7, 1.0, "web_app.py  —  Flask REST API",
        "gunicorn · Flask-Limiter (300 req/h) · SQLAlchemy · psycopg",
        color=(0.17, 0.14, 0.10), text_color=C_GOLD, fontsize=9)
    box(6.3, 2.7, 2.0, 1.0, "api/index.py\nVercel shim",
        "serverless fn", color=(0.10, 0.12, 0.18), text_color=C_BLUE, fontsize=8)
    box(8.6, 2.7, 2.4, 1.0, "Supabase Auth\nJWT validation",
        "3-tier fallback", color=(0.12, 0.15, 0.18), text_color=C_TEAL, fontsize=8)
    box(11.2, 2.7, 2.5, 1.0, "Vercel\nServerless\n(producción)",
        "vercel.json routing", color=(0.10, 0.12, 0.18), text_color=C_BLUE, fontsize=8)

    # ── FILA 5: Frontend ──────────────────────────────────────────────────
    ax.text(5.0, 2.35, "FRONTENDS / CLIENTES", ha="center", color=C_GOLD,
            fontsize=8, fontweight="bold", alpha=0.6)
    box(0.3, 1.1, 3.8, 1.0, "connect_talleres.html",
        "SPA · MapLibre GL · deck.gl · Chart.js · Tailwind CSS\nTickets · Vista ejecutiva · Mapa",
        color=(0.10, 0.16, 0.22), text_color=C_BLUE, fontsize=8)
    box(4.4, 1.1, 3.0, 1.0, "viewer_hex.py\n(Streamlit)",
        "Fleet Intelligence · pydeck\nConecta directo a DB", color=(0.12, 0.15, 0.18), text_color=C_TEAL, fontsize=8)
    box(7.7, 1.1, 2.4, 1.0, "viewer.py\nviewer_osm.py\n(Streamlit)",
        "Bubble / OSM map", color=(0.12, 0.15, 0.18), text_color=C_TEAL, fontsize=8)
    box(10.4, 1.1, 3.3, 1.0, "Browser / Usuario final",
        "HTTPS · Bearer token · Supabase login",
        color=(0.12, 0.15, 0.12), text_color=C_GREEN, fontsize=8)

    # ── FLECHAS ───────────────────────────────────────────────────────────
    # Fuentes → pipeline
    arrow(1.6, 7.6, 1.6, 6.85)
    arrow(4.6, 7.6, 3.5, 6.85)
    arrow(7.6, 7.6, 5.6, 6.85)
    arrow(10.6, 7.6, 5.8, 6.85)
    arrow(12.95, 7.6, 5.9, 6.85)
    # SAP cache → pipeline
    arrow(7.75, 6.35, 6.0, 6.35)
    # Pipeline → DB
    arrow(3.15, 5.85, 3.15, 5.3)
    # GH Actions → pipeline trigger
    arrow(9.3, 6.35, 6.0, 6.35)
    # DB → API
    arrow(3.15, 4.2, 3.15, 3.7)
    # API → Vercel shim
    arrow(6.0, 3.2, 6.3, 3.2)
    # Vercel shim → Vercel prod
    arrow(8.3, 3.2, 11.2, 3.2)
    # Supabase → API (auth)
    arrow(8.6, 3.2, 6.0, 3.2)
    # API → frontends
    arrow(3.15, 2.7, 3.15, 2.1)
    arrow(5.5, 2.7, 5.9, 2.1)
    # DB → viewer_hex (direct)
    arrow(5.5, 4.2, 5.9, 2.1)
    # Browser → Vercel
    arrow(12.05, 2.7, 12.05, 2.1)

    # ── LEYENDA ───────────────────────────────────────────────────────────
    legends = [
        (C_BLUE,   "APIs / Cloud"),
        (C_GOLD,   "Pipeline / Backend Python"),
        (C_GREEN,  "Datos / Storage local"),
        (C_TEAL,   "Streamlit / Auth"),
        (C_PURPLE, "ERP SAP / Cache"),
    ]
    for i, (c, lbl) in enumerate(legends):
        ax.add_patch(mpatches.Rectangle((0.3 + i * 2.7, 0.15), 0.3, 0.3, color=c))
        ax.text(0.65 + i * 2.7, 0.3, lbl, color=(0.75, 0.70, 0.65), fontsize=6.5, va="center")

    plt.tight_layout(pad=0.2)
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight",
                facecolor=C_DARK, edgecolor="none")
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────
# GENERAR DOCUMENTO WORD
# ─────────────────────────────────────────────
def build_doc():
    doc = Document()

    # ── Márgenes ──
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    styles = doc.styles

    # ── PORTADA ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GEOWORKSHOP")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0x8C)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Plataforma de Inteligencia de Cobertura de Flota")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Informe de Arquitectura y Diseño de Sistema\n{datetime.date.today().strftime('%d de %B de %Y')}")
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Grupo Kaufmann — Área de Tecnología")
    r4.bold = True
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ── 1. DESCRIPCIÓN GENERAL ───────────────────────────────────────────
    doc.add_heading("1. Descripción General del Sistema", level=1)
    doc.add_paragraph(
        "Geoworkshop es una plataforma de inteligencia geoespacial desarrollada para Grupo Kaufmann. "
        "Su propósito principal es determinar, de forma automática y diaria, qué vehículos de la flota "
        "se encuentran dentro del radio de cobertura de 100 km de cada taller (sucursal Kaufmann), "
        "facilitando la planificación de mantenimiento preventivo y correctivo."
    )
    doc.add_paragraph(
        "El sistema integra datos GPS en tiempo casi real desde múltiples fuentes (Copiloto y Geotab), "
        "los enriquece con información del ERP SAP de Kaufmann y con el maestro de flota en Excel, "
        "calcula distancias usando la fórmula de Haversine vectorizada, y persiste los resultados en una "
        "base de datos PostgreSQL/PostGIS en la nube (Neon). Los resultados son consumidos por una API "
        "REST Flask y visualizados en dashboards interactivos (HTML SPA y Streamlit)."
    )

    doc.add_heading("Capacidades principales", level=2)
    caps = [
        "Cobertura geoespacial: asignación de unidades a talleres por radio de 100 km (modos overlap y exclusive).",
        "Pipeline automatizado: ejecución diaria vía GitHub Actions a las 14:00 UTC (11:00 Santiago).",
        "Multi-fuente: integración simultánea de Copiloto (Chile/Perú/Paraguay) y Geotab (Colombia + Chile buses).",
        "Enriquecimiento ERP: modelo, marca, serie y segmento desde SAP Kaufmann por VIN.",
        "Decode de VIN: identificación de marca/modelo vía NHTSA API y tabla WMI como fallback.",
        "Gestión de tickets: módulo de mantenimiento con SLA configurable y notas por ticket.",
        "API REST segura: autenticación Bearer via Supabase, rate limiting (300 req/h), exportación CSV.",
        "Vista ejecutiva: resumen por zona/país, modelos × taller, KPIs y gráficos de cobertura.",
    ]
    for c in caps:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(c)

    # ── 2. DIAGRAMA DE FLUJO ─────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("2. Diagrama de Flujo y Arquitectura", level=1)
    doc.add_paragraph(
        "El siguiente diagrama muestra el flujo completo de datos desde las fuentes externas "
        "hasta los clientes finales, pasando por el pipeline de procesamiento, la base de datos "
        "y la capa de API."
    )

    img_buf = _draw_flow()
    doc.add_picture(img_buf, width=Inches(6.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_cap = doc.add_paragraph("Figura 1 — Arquitectura completa de Geoworkshop")
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.runs[0].italic = True
    p_cap.runs[0].font.size = Pt(9)
    p_cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ── 3. FLUJO DE DATOS DETALLADO ──────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. Flujo de Datos Detallado", level=1)

    steps = [
        ("Paso 1 — Autenticación y descarga GPS (Copiloto)",
         "app.py se autentica con la API de Copiloto usando un token estático (COPILOTO_API_TOKEN) "
         "o mediante email/contraseña como fallback. Descarga un CSV con todos los registros vehiculares "
         "(coordenadas GPS, IMEI, nombre de unidad, odómetro CAN, horometro)."),
        ("Paso 2 — Carga de datos Geotab (multi-base)",
         "Para cada base de datos configurada en GEOTAB_DATABASES (divemotor_colombia, divemotor, "
         "divemotor_buses), se conecta a la API de Geotab, obtiene el listado de dispositivos con su "
         "última posición GPS y odómetro/horometro CAN. Los VINs se decodifican via NHTSA o tabla WMI "
         "y se cachean en vin_cache.json. Las unidades Geotab se añaden solo si su unit_id no existe "
         "ya en el dataset Copiloto (Copiloto tiene precedencia)."),
        ("Paso 3 — Enriquecimiento con Master de Flota y SAP",
         "Se cruzan las unidades por VIN (primario) e IMEI (fallback) contra master_Flota.xlsx para "
         "obtener Empresa, Marca, Modelo y Patente. Luego se aplica el cache SAP (sap_vin_cache.json) "
         "que normaliza modelo, marca, serie, segmento, automotora y RUT cliente desde el ERP Kaufmann. "
         "El cache SAP se construye por separado con build_sap_cache.py."),
        ("Paso 4 — Cálculo de distancias (Haversine vectorizado)",
         "Se calcula la distancia en km desde cada unidad GPS a cada taller usando Haversine vectorizado "
         "con NumPy (sin PostGIS, para máxima velocidad). Se aplican dos modos: overlap (todos los talleres "
         "en el radio) y exclusive (solo el más cercano). El radio por defecto es 100 km (RADIUS_KM)."),
        ("Paso 5 — Persistencia en PostgreSQL/Neon",
         "Los resultados se escriben en cuatro tablas snapshot: snapshot_run (metadatos del run), "
         "snapshot_unit (una fila por vehículo con coordenadas, empresa, taller asignado y distancia), "
         "snapshot_taller_overlap y snapshot_taller_exclusive (conteos por taller). La tabla dim_taller "
         "se hace upsert con los datos maestros de cada taller. Se genera también un CSV local en Scripts/out/."),
        ("Paso 6 — API REST (web_app.py)",
         "Flask lee solo desde la DB, nunca llama a Copiloto. Cada request resuelve el run más reciente "
         "via _latest_run(). La autenticación Bearer usa Supabase con fallback a JWT local. "
         "Los endpoints sirven KPIs, tabla ejecutiva, detalle de unidades, tendencia semanal, "
         "modelos × taller y gestión de tickets de mantenimiento."),
        ("Paso 7 — Entrega al usuario final",
         "En producción la API corre como función serverless en Vercel (api/index.py como shim). "
         "El SPA connect_talleres.html se sirve en la ruta raíz y consume los endpoints /api/*. "
         "Los dashboards Streamlit (viewer_hex.py, viewer.py, viewer_osm.py) se conectan directamente "
         "a la DB via SQLAlchemy para mayor rendimiento."),
    ]
    for title, body in steps:
        h = doc.add_heading(title, level=2)
        doc.add_paragraph(body)

    # ── 4. SERVICIOS / APIs CONSUMIDOS ───────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. Servicios y APIs Externos Consumidos", level=1)

    svc_data = [
        ("Copiloto API",           "https://api.copiloto.ai",                       "REST / CSV",    "GPS, IMEI, odómetro, horometro de flota Chile/Perú/Paraguay",       "COPILOTO_API_TOKEN\no COPILOTO_EMAIL\n+ COPILOTO_PASSWORD"),
        ("Copiloto Sign-in",       "https://accounts.copiloto.ai",                  "REST / JSON",   "Autenticación cuando no hay token estático",                         "COPILOTO_EMAIL\nCOPILOTO_PASSWORD"),
        ("Geotab API",             "https://my.geotab.com",                          "JSON-RPC",      "GPS, odómetro, VIN para 3 bases de datos (CO, CL, buses)",           "GEOTAB_USERNAME\nGEOTAB_PASSWORD\nGEOTAB_DATABASES"),
        ("SAP ERP Kaufmann",       "https://apimaz.grupokaufmann.com/prd/erp/…",    "REST / JSON",   "Marca, modelo, serie, segmento, automotora por VIN",                 "ERP_SUBSCRIPTION_KEY"),
        ("NHTSA vPIC API",         "https://vpic.nhtsa.dot.gov/api",                "REST / JSON",   "Decode de VIN norteamericanos (prefijo 1/2/3) → marca + modelo",    "(pública, sin clave)"),
        ("PostgreSQL / Neon",      "ep-old-sky-*.neon.tech:5432",                   "psycopg3",      "Almacenamiento de snapshots, dim_taller, tickets",                   "DATABASE_URL\no PGHOST/USER/PASS"),
        ("Supabase Auth",          "https://<project>.supabase.co/auth/v1/user",    "REST / JWT",    "Validación de tokens Bearer de usuarios frontend",                   "SUPABASE_URL\nSUPABASE_ANON_KEY\nSUPABASE_JWT_SECRET"),
        ("GitHub Actions",         "github.com / Actions runner",                   "YAML / cron",   "Ejecución diaria del pipeline app.py",                               "Secrets del repositorio"),
        ("Vercel Serverless",      "vercel.com",                                     "Serverless fn", "Hosting de la API Flask y el SPA HTML en producción",                "Variables en dashboard Vercel"),
    ]

    tbl = doc.add_table(rows=1 + len(svc_data), cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _heading_row(tbl, "Servicio", "Endpoint / Host", "Protocolo", "Datos que provee", "Variables de entorno")
    _set_col_widths(tbl, [3.0, 3.5, 1.8, 5.0, 3.0])

    for i, (svc, ep, proto, datos, envs) in enumerate(svc_data):
        row = tbl.rows[i + 1]
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _cell_bg(cell, bg)
        row.cells[0].paragraphs[0].add_run(svc).bold = True
        row.cells[1].paragraphs[0].add_run(ep).font.size = Pt(8)
        row.cells[2].paragraphs[0].add_run(proto).font.size = Pt(8)
        row.cells[3].paragraphs[0].add_run(datos).font.size = Pt(8)
        r_env = row.cells[4].paragraphs[0].add_run(envs)
        r_env.font.size = Pt(8)
        r_env.font.name = "Courier New"

    # ── 5. TECNOLOGÍAS Y LENGUAJES ───────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("5. Tecnologías, Lenguajes y Librerías", level=1)

    tech_data = [
        ("Python 3.11",        "Pipeline (app.py), API (web_app.py), scripts auxiliares",        "Lenguaje principal del backend"),
        ("pandas / NumPy",     "app.py, web_app.py",                                             "Manipulación de DataFrames y cálculo Haversine vectorizado"),
        ("SQLAlchemy 2.x",     "app.py, web_app.py, viewer_hex.py",                              "ORM / query builder; conexión a PostgreSQL via psycopg3"),
        ("psycopg3",           "app.py, web_app.py",                                             "Driver PostgreSQL nativo para Python (psycopg[binary])"),
        ("Flask",              "web_app.py",                                                     "Framework web ligero para la API REST"),
        ("Flask-Limiter",      "web_app.py",                                                     "Rate limiting: 300 req/h global, 10/min en /api/export"),
        ("gunicorn",           "Despliegue local / Docker",                                      "Servidor WSGI para producción local"),
        ("Streamlit",          "viewer_hex.py, viewer.py, viewer_osm.py",                        "Dashboards interactivos de visualización"),
        ("pydeck",             "viewer_hex.py",                                                  "Capas hexagonales 3D sobre mapa (deck.gl desde Python)"),
        ("requests",           "app.py, build_sap_cache.py",                                     "Llamadas HTTP a Copiloto, Geotab, SAP, NHTSA"),
        ("openpyxl",           "app.py, build_sap_cache.py",                                     "Lectura de master_Flota.xlsx y talleres.xlsx"),
        ("python-dotenv",      "app.py, web_app.py, build_sap_cache.py",                         "Carga de variables de entorno desde .env"),
        ("PyJWT",              "web_app.py",                                                     "Decodificación y verificación de tokens JWT (Supabase)"),
        ("JavaScript (ES2020)","connect_talleres.html",                                          "Lógica del SPA: filtros, tablas, auth, fetch a /api/*"),
        ("MapLibre GL JS",     "connect_talleres.html",                                          "Mapa vectorial interactivo (vía CDN)"),
        ("deck.gl",            "connect_talleres.html",                                          "Capas de visualización geoespacial sobre el mapa"),
        ("Chart.js",           "connect_talleres.html",                                          "Gráficos de barras, donut y series temporales"),
        ("Tailwind CSS",       "connect_talleres.html",                                          "Estilos utility-first (vía CDN)"),
        ("PostgreSQL 15+",     "Neon (producción), Docker (desarrollo local)",                   "Base de datos relacional con extensión PostGIS"),
        ("PostGIS",            "geo-workshop-db/init/001_init.sql",                              "Extensión geoespacial habilitada; geometría en dim_taller"),
        ("GitHub Actions",     ".github/workflows/daily_pipeline.yml",                           "CI/CD: ejecución diaria del pipeline, gestión de secrets"),
        ("Vercel",             "api/index.py + vercel.json",                                     "Hosting serverless de la API Flask y el SPA en producción"),
        ("Supabase",           "Proyecto cloud",                                                 "Autenticación de usuarios (JWT Bearer)"),
        ("Neon",               "PostgreSQL serverless en AWS us-east-1",                         "Base de datos de producción con connection pooling"),
        ("Docker Compose",     "geo-workshop-db/",                                               "PostgreSQL + PostGIS local para desarrollo"),
    ]

    tbl2 = doc.add_table(rows=1 + len(tech_data), cols=3)
    tbl2.style = "Table Grid"
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _heading_row(tbl2, "Tecnología / Librería", "Dónde se usa", "Función")
    _set_col_widths(tbl2, [4.0, 5.0, 7.2])

    for i, (tech, donde, funcion) in enumerate(tech_data):
        row = tbl2.rows[i + 1]
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _cell_bg(cell, bg)
        r = row.cells[0].paragraphs[0].add_run(tech)
        r.bold = True; r.font.size = Pt(8)
        row.cells[1].paragraphs[0].add_run(donde).font.size = Pt(8)
        row.cells[2].paragraphs[0].add_run(funcion).font.size = Pt(8)

    # ── 6. ENDPOINTS API REST ────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("6. Endpoints de la API REST", level=1)
    doc.add_paragraph(
        "Todos los endpoints requieren Authorization: Bearer <supabase_token> en producción. "
        "URL base en producción: https://<proyecto>.vercel.app"
    )

    ep_data = [
        ("GET",   "/",                             "SPA HTML",         "Sirve connect_talleres.html (sin caché)"),
        ("GET",   "/api/data",                     "JSON",             "Marcadores de mapa + KPIs generales"),
        ("GET",   "/api/ejecutivo",                "JSON",             "Tabla resumen por taller, zonas, cobertura, top distancias"),
        ("GET",   "/api/detalle",                  "JSON (≤500 rows)", "Listado de unidades; filtros: taller, empresa, q, max_dist"),
        ("GET",   "/api/tendencia",                "JSON",             "Serie temporal semanal de todos los runs históricos"),
        ("GET",   "/api/modelos-sucursal",         "JSON",             "Matriz de modelos por taller (para tabla × taller)"),
        ("GET",   "/api/radio-search",             "JSON",             "Búsqueda ad-hoc: ?lat=&lon=&radius_km="),
        ("GET",   "/api/estado-flota",             "JSON",             "Estado OBD: odómetro vs umbral por marca + fallas DTC"),
        ("GET",   "/api/talleres",                 "JSON",             "Lista de talleres activos desde dim_taller"),
        ("GET",   "/api/export/<tipo>",            "CSV",              "Descarga completa: tipo = units | detalle | cobertura"),
        ("GET",   "/api/tickets",                  "JSON",             "Lista tickets; filtros: unit_id, estado=vencido"),
        ("POST",  "/api/tickets",                  "JSON",             "Crear nuevo ticket de mantenimiento"),
        ("GET",   "/api/tickets/<id>",             "JSON",             "Detalle de un ticket"),
        ("PATCH", "/api/tickets/<id>",             "JSON",             "Actualizar estado, prioridad, asignación"),
        ("POST",  "/api/tickets/<id>/notes",       "JSON",             "Añadir nota a un ticket"),
        ("GET",   "/api/tickets/kpis",             "JSON",             "KPIs agregados: conteos por estado, vencidos"),
        ("GET",   "/api/maintenance-records",      "JSON",             "Lista de registros de trabajo completado"),
        ("POST",  "/api/maintenance-records",      "JSON",             "Crear registro de mantenimiento completado"),
    ]

    tbl3 = doc.add_table(rows=1 + len(ep_data), cols=4)
    tbl3.style = "Table Grid"
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _heading_row(tbl3, "Método", "Ruta", "Respuesta", "Descripción")
    _set_col_widths(tbl3, [1.5, 4.5, 2.5, 7.7])

    METHOD_COLOR = {
        "GET":   ("E8F4FD", "1A568C"),
        "POST":  ("E8F9EE", "1A6B3C"),
        "PATCH": ("FFF8E8", "8C6B1A"),
    }
    for i, (method, route, resp, desc) in enumerate(ep_data):
        row = tbl3.rows[i + 1]
        bg_row = "F8F8F8" if i % 2 == 0 else "FFFFFF"
        m_bg, m_fg = METHOD_COLOR.get(method, ("F0F0F0", "333333"))
        _cell_bg(row.cells[0], m_bg)
        for j in range(1, 4):
            _cell_bg(row.cells[j], bg_row)
        r_m = row.cells[0].paragraphs[0].add_run(method)
        r_m.bold = True; r_m.font.size = Pt(8)
        r_m.font.color.rgb = RGBColor.from_string(m_fg)
        r_r = row.cells[1].paragraphs[0].add_run(route)
        r_r.font.name = "Courier New"; r_r.font.size = Pt(8)
        row.cells[2].paragraphs[0].add_run(resp).font.size = Pt(8)
        row.cells[3].paragraphs[0].add_run(desc).font.size = Pt(8)

    # ── 7. ESQUEMA DE BASE DE DATOS ──────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("7. Esquema de Base de Datos", level=1)
    doc.add_paragraph(
        "PostgreSQL con extensión PostGIS. El esquema se crea automáticamente en el primer run "
        "de app.py mediante run_migrations(). Las tablas snapshot son append-only (una fila por "
        "run_id); dim_taller se actualiza por upsert."
    )

    db_tables = [
        ("snapshot_run",              "run_id, timestamp, unit_count, config params",
         "Metadatos de cada ejecución del pipeline. Clave primaria run_id (UUID)."),
        ("snapshot_unit",             "run_id, unit_id (VIN/IMEI), lat, lon, empresa, modelo, marca, patente,\ntaller asignado, distancia km, dentro_radio, source, sap_*",
         "Una fila por vehículo por run. Contiene coordenadas GPS, info de empresa y asignación al taller más cercano."),
        ("snapshot_taller_overlap",   "run_id, taller_cercano_nombre, unidades",
         "Conteo de unidades dentro del radio para cada taller (modo overlap: todas las unidades en radio, sin exclusividad)."),
        ("snapshot_taller_exclusive", "run_id, taller_cercano_nombre, unidades",
         "Conteo solo con el taller más cercano por unidad (modo exclusive)."),
        ("dim_taller",                "id, nombre, lat, lon, zona, pais, geometry (PostGIS)",
         "Dimensión estática de talleres. Upsert en cada pipeline run. La columna geometry permite consultas PostGIS."),
        ("maintenance_ticket",        "id, unit_id, estado, prioridad, assigned_to, created_by,\ncreated_at, updated_at, sla_days",
         "Tickets de mantenimiento. Estado: pendiente, en_proceso, completado, cerrado, cancelado. SLA en días hábiles."),
        ("maintenance_ticket_note",   "id, ticket_id (FK), autor, texto, created_at",
         "Notas libres adjuntas a un ticket. FK a maintenance_ticket."),
        ("maintenance_record",        "id, unit_id, tipo, descripcion, fecha, tecnico, ticket_id (FK opt.)",
         "Registro de trabajo efectivamente realizado. Puede vincularse a un ticket o crearse de forma independiente."),
    ]

    tbl4 = doc.add_table(rows=1 + len(db_tables), cols=3)
    tbl4.style = "Table Grid"
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    _heading_row(tbl4, "Tabla", "Columnas principales", "Descripción")
    _set_col_widths(tbl4, [4.0, 6.5, 5.7])

    for i, (tname, cols, desc) in enumerate(db_tables):
        row = tbl4.rows[i + 1]
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _cell_bg(cell, bg)
        r = row.cells[0].paragraphs[0].add_run(tname)
        r.bold = True; r.font.name = "Courier New"; r.font.size = Pt(8)
        r2 = row.cells[1].paragraphs[0].add_run(cols)
        r2.font.name = "Courier New"; r2.font.size = Pt(7.5)
        row.cells[2].paragraphs[0].add_run(desc).font.size = Pt(8)

    # ── 8. AUTOMATIZACIÓN ────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("8. Automatización y Despliegue", level=1)

    doc.add_heading("GitHub Actions — Pipeline diario", level=2)
    doc.add_paragraph(
        "El archivo .github/workflows/daily_pipeline.yml ejecuta app.py cada día a las 14:00 UTC "
        "(11:00 AM Santiago). También puede dispararse manualmente (workflow_dispatch). "
        "Las credenciales sensibles (contraseñas de Copiloto, Geotab, Neon y la clave del ERP SAP) "
        "se almacenan como secrets del repositorio GitHub. El host, base de datos y usuario de Neon "
        "están hardcodeados en el YAML."
    )

    doc.add_heading("Vercel — Producción", level=2)
    doc.add_paragraph(
        "La API Flask se despliega como función serverless en Vercel a través del shim api/index.py, "
        "que añade Scripts/ al sys.path de Python y re-exporta el objeto app de Flask. "
        "vercel.json enruta todo el tráfico a esa función. "
        "El HTML del SPA se sirve en la ruta raíz con header Cache-Control: no-store "
        "para garantizar que los usuarios siempre reciban la versión más reciente. "
        "Las variables de entorno deben configurarse en el dashboard de Vercel."
    )

    doc.add_heading("Build SAP cache — Ejecución manual", level=2)
    doc.add_paragraph(
        "build_sap_cache.py consulta el ERP SAP de Kaufmann por cada VIN único en master_Flota.xlsx "
        "y guarda los resultados en Data/sap_vin_cache.json. Este script se ejecuta manualmente "
        "(o como paso previo al pipeline cuando hay VINs nuevos) y no forma parte del cron diario. "
        "Durante el pipeline diario, app.py solo lee el cache — nunca llama directamente al SAP."
    )

    # ── PIE DE PÁGINA ────────────────────────────────────────────────────
    section = doc.sections[0]
    footer  = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run(
        f"Geoworkshop — Informe de Arquitectura | Grupo Kaufmann | {datetime.date.today().year}"
    ).font.size = Pt(8)

    out_path = "Geoworkshop_Arquitectura.docx"
    doc.save(out_path)
    print(f"Documento generado: {out_path}")


if __name__ == "__main__":
    build_doc()
