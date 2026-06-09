# -*- coding: utf-8 -*-
"""
Genera Geoworkshop_Infraestructura.docx con la arquitectura completa
para despliegue en AWS (RDS) + servidor externo a Vercel + Azure AD.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color="1F3864"):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def table_header_row(table, headers, bg="1F3864", fg="FFFFFF"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(fg)
                run.font.size = Pt(10)

def add_table_row(table, values, bold_first=False, bg=None, center=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        if bg:
            set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            if center:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(9)
                if bold_first and i == 0:
                    run.bold = True

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def add_note(doc, text, color="FFF3CD"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run("   " + text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x7D, 0x60, 0x08)

# ── Documento ────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Estilo base
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("GEOWORKSHOP")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Arquitectura e Infraestructura de Producción")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(
    f"AWS RDS · Servidor Dedicado · Azure AD\n"
    f"Grupo Kaufmann — {datetime.date.today().strftime('%B %Y')}"
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Resumen Ejecutivo", level=1)
doc.add_paragraph(
    "Geoworkshop es una plataforma de inteligencia de cobertura de flota para Grupo Kaufmann. "
    "Consume datos GPS desde Copiloto API y Geotab, calcula qué vehículos se encuentran dentro "
    "del radio de cobertura de cada taller (100 km por defecto), y expone ese análisis a través "
    "de una API Flask y un SPA (Single Page Application) HTML."
)
doc.add_paragraph(
    "Este documento define la arquitectura de producción con tres componentes de infraestructura principales:"
)
items = [
    "Base de datos PostgreSQL + PostGIS alojada en AWS RDS.",
    "Servidor de aplicación (Flask API + pipeline diario) en un servidor Linux dedicado fuera de Vercel.",
    "Autenticación de usuarios mediante Azure Active Directory (Azure AD / Microsoft Entra ID).",
]
for item in items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(10)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA GENERAL
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Arquitectura General", level=1)

doc.add_paragraph(
    "La arquitectura sigue un modelo de tres capas: autenticación (Azure AD), "
    "capa de aplicación (servidor Linux), y capa de datos (AWS RDS). "
    "El pipeline de datos se ejecuta diariamente mediante GitHub Actions o cron job."
)

diagram = (
    "┌─────────────────────────────────────────────────────────────────────┐\n"
    "│                        FUENTES DE DATOS                             │\n"
    "│          Copiloto API (GPS CSV)  ·  Geotab API  ·  SAP ERP          │\n"
    "└──────────────────────────┬──────────────────────────────────────────┘\n"
    "                           │  app.py  (pipeline diario · cron 14:00 UTC)\n"
    "                           ▼\n"
    "┌─────────────────────────────────────────────────────────────────────┐\n"
    "│                    AWS RDS — PostgreSQL 15 + PostGIS                 │\n"
    "│             Subnet privada · cifrado en reposo (AES-256)             │\n"
    "└──────────┬────────────────────────────────────────────────────────┬─┘\n"
    "           │  SQLAlchemy (psycopg)                                  │\n"
    "           ▼                                                        ▼\n"
    "┌──────────────────────────┐                           ┌────────────────────┐\n"
    "│  Servidor de Aplicación  │                           │  viewer_hex.py     │\n"
    "│  Ubuntu 22.04 LTS        │                           │  Streamlit UI      │\n"
    "│  nginx (443/TLS)         │                           │  (opcional · mismo │\n"
    "│  gunicorn → web_app.py   │                           │   servidor · :8501)│\n"
    "│  cron → app.py           │                           └────────────────────┘\n"
    "└──────────┬───────────────┘\n"
    "           │  /api/*  (JSON)\n"
    "           ▼\n"
    "┌─────────────────────────────────────────────────────────────────────┐\n"
    "│              connect_talleres.html  (SPA · Browser del usuario)      │\n"
    "│         MapLibre GL · deck.gl · Chart.js · Tailwind CSS (CDN)        │\n"
    "│                    MSAL.js → Azure AD (OAuth 2.0)                    │\n"
    "└────────────────────────────┬────────────────────────────────────────┘\n"
    "                             │  Bearer JWT (RS256)\n"
    "                             ▼\n"
    "┌─────────────────────────────────────────────────────────────────────┐\n"
    "│                  Microsoft Azure AD / Entra ID                       │\n"
    "│            App Registration · JWKS endpoint · Tenant ID              │\n"
    "└─────────────────────────────────────────────────────────────────────┘"
)
add_code_block(doc, diagram)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 3. ESPECIFICACIONES DE HARDWARE
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Especificaciones de Hardware (Máquinas Virtuales)", level=1)

doc.add_paragraph(
    "Se definen tres ambientes: Desarrollo, Staging y Producción. "
    "Los tamaños están dimensionados para una flota de hasta 5.000 vehículos con "
    "consultas concurrentes de hasta 50 usuarios simultáneos."
)

# ── 3.1 Servidor de Aplicación ──
heading(doc, "3.1 Servidor de Aplicación (Flask API + Pipeline)", level=2)

t = doc.add_table(rows=1, cols=5)
t.style = "Table Grid"
t.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t, ["Ambiente", "Tipo / Instancia", "vCPU", "RAM", "Almacenamiento"])

env_specs = [
    ("Desarrollo",  "EC2 t3.small  \n(o equivalente)",   "2", "2 GB",  "30 GB SSD gp3"),
    ("Staging",     "EC2 t3.medium \n(o equivalente)",   "2", "4 GB",  "50 GB SSD gp3"),
    ("Producción",  "EC2 t3.large  \n(o equivalente)",   "2", "8 GB",  "80 GB SSD gp3"),
]
row_colors = ["F2F2F2", "FFFFFF", "F2F2F2"]
for (env, tipo, vcpu, ram, disk), bg in zip(env_specs, row_colors):
    row = t.add_row()
    vals = [env, tipo, vcpu, ram, disk]
    for i, v in enumerate(vals):
        row.cells[i].text = v
        set_cell_bg(row.cells[i], bg)
        for para in row.cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
add_note(doc,
    "En producción se recomienda t3.large. El pipeline diario (app.py) puede "
    "consumir hasta 1.5 GB de RAM durante la carga y enriquecimiento de datos. "
    "gunicorn corre con 4 workers (≈200 MB c/u).")
doc.add_paragraph()

# ── 3.2 Base de Datos AWS RDS ──
heading(doc, "3.2 Base de Datos — AWS RDS PostgreSQL", level=2)

t2 = doc.add_table(rows=1, cols=6)
t2.style = "Table Grid"
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t2, ["Ambiente", "Instancia RDS", "vCPU", "RAM", "Almacenamiento", "Multi-AZ"])

rds_specs = [
    ("Desarrollo",  "db.t3.micro",   "2", "1 GB",  "20 GB gp3",  "No"),
    ("Staging",     "db.t3.small",   "2", "2 GB",  "50 GB gp3",  "No"),
    ("Producción",  "db.t3.medium",  "2", "4 GB",  "100 GB gp3", "Sí (recomendado)"),
]
row_colors2 = ["F2F2F2", "FFFFFF", "F2F2F2"]
for (env, inst, vcpu, ram, disk, maz), bg in zip(rds_specs, row_colors2):
    row = t2.add_row()
    vals = [env, inst, vcpu, ram, disk, maz]
    for i, v in enumerate(vals):
        row.cells[i].text = v
        set_cell_bg(row.cells[i], bg)
        for para in row.cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

doc.add_paragraph()
add_note(doc,
    "El crecimiento estimado de datos es ~500 MB por año (snapshot diario de 5.000 unidades). "
    "Habilitar autoscaling de almacenamiento en RDS como precaución. "
    "Backup automático: retención de 7 días mínimo.")
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 4. CONFIGURACIÓN AWS
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Configuración AWS", level=1)

# ── 4.1 Red (VPC) ──
heading(doc, "4.1 Red — VPC y Subnets", level=2)
doc.add_paragraph(
    "Se requiere una VPC dedicada con subnets públicas y privadas en al menos una "
    "zona de disponibilidad (dos para producción con Multi-AZ)."
)

net_items = [
    "VPC: 10.0.0.0/16",
    "Subnet pública (EC2 app server): 10.0.1.0/24  —  az-a",
    "Subnet privada (RDS): 10.0.2.0/24  —  az-a",
    "Subnet privada (RDS Multi-AZ réplica): 10.0.3.0/24  —  az-b  (solo producción)",
    "Internet Gateway: asociado a la subnet pública",
    "NAT Gateway: para que el EC2 acceda a APIs externas (Copiloto, Geotab, Azure) sin IP pública directa (opcional pero recomendado en producción)",
]
for item in net_items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── 4.2 Security Groups ──
heading(doc, "4.2 Security Groups", level=2)

t3 = doc.add_table(rows=1, cols=4)
t3.style = "Table Grid"
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t3, ["Recurso", "Puerto", "Protocolo", "Origen permitido"])

sg_rules = [
    ("EC2 App Server",  "443 (HTTPS)",  "TCP", "0.0.0.0/0 (internet)"),
    ("EC2 App Server",  "80 (HTTP)",    "TCP", "0.0.0.0/0 → redirige a 443"),
    ("EC2 App Server",  "22 (SSH)",     "TCP", "IP del equipo de desarrollo únicamente"),
    ("EC2 App Server",  "8501 (Streamlit)", "TCP", "IP interna / VPN (opcional)"),
    ("RDS PostgreSQL",  "5432",         "TCP", "Security Group del EC2 únicamente"),
]
for i, (res, port, proto, src) in enumerate(sg_rules):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    add_table_row(t3, [res, port, proto, src], bold_first=True, bg=bg)

doc.add_paragraph()

# ── 4.3 RDS Setup ──
heading(doc, "4.3 Configuración RDS", level=2)
rds_steps = [
    "Engine: PostgreSQL 15.x",
    "Habilitar PostGIS después del primer arranque:  CREATE EXTENSION IF NOT EXISTS postgis;",
    "DB name: geocobertura  (o el que se prefiera)",
    "Cifrado en reposo: activado (KMS key administrada por AWS)",
    "Cifrado en tránsito: SSL requerido (sslmode=require en DATABASE_URL)",
    "Backup automático: 7 días de retención, ventana de mantenimiento fuera de horario de pipeline (no entre 13:00-15:00 UTC)",
    "Parameter group: default.postgres15 (sin cambios necesarios para este proyecto)",
]
for step in rds_steps:
    p = doc.add_paragraph(step, style="List Bullet")
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()
add_code_block(doc,
    "# DATABASE_URL resultante\n"
    "DATABASE_URL=postgresql+psycopg://geo_user:PASSWORD@"
    "<rds-endpoint>.rds.amazonaws.com:5432/geocobertura?sslmode=require"
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 5. CONFIGURACIÓN AZURE AD
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Configuración Azure AD (Microsoft Entra ID)", level=1)

doc.add_paragraph(
    "Azure AD reemplaza a Supabase como proveedor de identidad. "
    "El flujo es OAuth 2.0 Authorization Code con PKCE. "
    "El frontend obtiene un JWT firmado con RS256; el backend lo valida "
    "contra el JWKS endpoint público de Azure."
)

# ── 5.1 App Registration ──
heading(doc, "5.1 App Registration en Azure Portal", level=2)
az_steps = [
    "Azure Portal → Microsoft Entra ID → App registrations → New registration",
    "Name: Geoworkshop  |  Supported account types: Single tenant (solo Grupo Kaufmann)",
    "Redirect URI: https://<dominio-del-servidor>/  (tipo Single-page application)",
    "En 'Authentication': habilitar Access tokens y ID tokens (implicit grant + hybrid flows)",
    "En 'API permissions': añadir Microsoft Graph → User.Read (delegated)",
    "En 'Expose an API': crear scope  api://<CLIENT_ID>/access  (para validación de audience)",
    "Copiar: Application (client) ID y Directory (tenant) ID — se usan en las variables de entorno",
]
for s in az_steps:
    p = doc.add_paragraph(s, style="List Number")
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── 5.2 Variables de entorno Azure ──
heading(doc, "5.2 Variables de entorno requeridas", level=2)

t4 = doc.add_table(rows=1, cols=3)
t4.style = "Table Grid"
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t4, ["Variable", "Valor", "Dónde se obtiene"])

az_vars = [
    ("AZURE_TENANT_ID",  "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx", "Azure Portal → Entra ID → Overview → Directory ID"),
    ("AZURE_CLIENT_ID",  "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx", "App Registration → Overview → Application ID"),
    ("AZURE_CLIENT_SECRET", "generado en portal", "App Registration → Certificates & secrets → New secret (solo si hay flujo server-to-server)"),
]
for i, row_data in enumerate(az_vars):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    add_table_row(t4, list(row_data), bold_first=True, bg=bg)

doc.add_paragraph()

# ── 5.3 Cambios en el código ──
heading(doc, "5.3 Cambios requeridos en el código", level=2)

doc.add_paragraph("Backend — web_app.py: reemplazar _verify_supabase_token por validación RS256:")
add_code_block(doc,
    "# pip install PyJWT cryptography requests\n"
    "import jwt, requests\n\n"
    "AZURE_TENANT_ID = os.getenv('AZURE_TENANT_ID')\n"
    "AZURE_CLIENT_ID = os.getenv('AZURE_CLIENT_ID')\n"
    "_JWKS_CACHE = {}\n\n"
    "def _verify_azure_token(token: str) -> bool:\n"
    "    jwks_url = (f'https://login.microsoftonline.com/'\n"
    "               f'{AZURE_TENANT_ID}/discovery/v2.0/keys')\n"
    "    if not _JWKS_CACHE:\n"
    "        _JWKS_CACHE.update(requests.get(jwks_url, timeout=5).json())\n"
    "    header = jwt.get_unverified_header(token)\n"
    "    key_data = next(k for k in _JWKS_CACHE['keys']\n"
    "                    if k['kid'] == header['kid'])\n"
    "    pub_key = jwt.algorithms.RSAAlgorithm.from_jwk(key_data)\n"
    "    jwt.decode(token, pub_key, algorithms=['RS256'],\n"
    "               audience=AZURE_CLIENT_ID)\n"
    "    return True  # lanza excepción si es inválido"
)
doc.add_paragraph()

doc.add_paragraph("Frontend — connect_talleres.html: reemplazar Supabase JS por MSAL.js:")
add_code_block(doc,
    "<!-- CDN MSAL.js -->\n"
    "<script src='https://alcdn.msauth.net/browser/2.38.0/js/msal-browser.min.js'></script>\n\n"
    "const msalConfig = {\n"
    "  auth: {\n"
    "    clientId: 'TU_AZURE_CLIENT_ID',\n"
    "    authority: 'https://login.microsoftonline.com/TU_TENANT_ID',\n"
    "    redirectUri: window.location.origin\n"
    "  }\n"
    "};\n"
    "const msal = new msal.PublicClientApplication(msalConfig);\n\n"
    "// En authFetch(): obtener token con\n"
    "// msal.acquireTokenSilent({ scopes: ['api://CLIENT_ID/access'] })\n"
    "// e inyectar como Authorization: Bearer <token>"
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 6. CONFIGURACIÓN DEL SERVIDOR DE APLICACIÓN
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Configuración del Servidor de Aplicación", level=1)

heading(doc, "6.1 Stack de software", level=2)
sw_items = [
    "OS: Ubuntu 22.04 LTS",
    "Python: 3.11  (instalar con deadsnakes PPA si no viene por defecto)",
    "nginx: 1.24+  (reverse proxy + SSL termination)",
    "gunicorn: 21+  (WSGI server para Flask)",
    "supervisor: proceso de control para gunicorn (reinicio automático)",
    "certbot: certificado TLS con Let's Encrypt (o ACM si hay ALB de AWS)",
    "cron: para el pipeline diario (o GitHub Actions, ya configurado)",
]
for item in sw_items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()

heading(doc, "6.2 Comandos de instalación y arranque", level=2)
add_code_block(doc,
    "# 1. Clonar repositorio\n"
    "git clone https://github.com/ConnectDev-ai/geoworkshop-kf.git /opt/geoworkshop\n"
    "cd /opt/geoworkshop\n\n"
    "# 2. Instalar dependencias Python\n"
    "pip install -r Scripts/requirements.txt\n\n"
    "# 3. Crear archivo de entorno\n"
    "cp .env.example .env   # editar con valores reales\n\n"
    "# 4. Ejecutar pipeline una vez (poblar DB)\n"
    "python Scripts/app.py\n\n"
    "# 5. Levantar API con gunicorn (supervisor lo mantiene vivo)\n"
    "gunicorn -w 4 -b 127.0.0.1:5000 web_app:app\n\n"
    "# 6. (Opcional) Streamlit en puerto 8501\n"
    "streamlit run Scripts/viewer_hex.py --server.port 8501 --server.address 127.0.0.1"
)
doc.add_paragraph()

heading(doc, "6.3 Configuración nginx (HTTPS)", level=2)
add_code_block(doc,
    "# /etc/nginx/sites-available/geoworkshop\n"
    "server {\n"
    "    listen 443 ssl;\n"
    "    server_name geoworkshop.grupokaufmann.com;\n\n"
    "    ssl_certificate     /etc/letsencrypt/live/.../fullchain.pem;\n"
    "    ssl_certificate_key /etc/letsencrypt/live/.../privkey.pem;\n\n"
    "    location / {\n"
    "        proxy_pass         http://127.0.0.1:5000;\n"
    "        proxy_set_header   Host $host;\n"
    "        proxy_set_header   X-Real-IP $remote_addr;\n"
    "        proxy_read_timeout 120s;\n"
    "    }\n"
    "}\n"
    "server {\n"
    "    listen 80;\n"
    "    server_name geoworkshop.grupokaufmann.com;\n"
    "    return 301 https://$host$request_uri;\n"
    "}"
)
doc.add_paragraph()

heading(doc, "6.4 Cron job — Pipeline diario", level=2)
add_code_block(doc,
    "# crontab -e  (usuario que corre la app)\n"
    "# Pipeline diario a las 14:00 UTC (11:00 Santiago)\n"
    "0 14 * * * cd /opt/geoworkshop && python Scripts/app.py >> /var/log/geoworkshop/pipeline.log 2>&1\n\n"
    "# Rotación de logs (logrotate)\n"
    "# /etc/logrotate.d/geoworkshop\n"
    "/var/log/geoworkshop/*.log { daily rotate 14 compress missingok }"
)
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 7. VARIABLES DE ENTORNO COMPLETAS
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Variables de Entorno Completas (.env)", level=1)

t5 = doc.add_table(rows=1, cols=3)
t5.style = "Table Grid"
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t5, ["Variable", "Valor de ejemplo / Descripción", "Requerida"])

env_vars = [
    ("DATABASE_URL",         "postgresql+psycopg://user:pass@<rds>.amazonaws.com:5432/geocobertura?sslmode=require", "Sí"),
    ("COPILOTO_API_TOKEN",   "token estático de Copiloto (preferido sobre email/password)",                          "Sí"),
    ("COPILOTO_EMAIL",       "email de Copiloto (fallback si no hay API token)",                                     "Condicional"),
    ("COPILOTO_PASSWORD",    "contraseña de Copiloto",                                                               "Condicional"),
    ("GEOTAB_DATABASES",     "divemotor_colombia,divemotor,divemotor_buses",                                         "Sí"),
    ("GEOTAB_USERNAME",      "usuario Geotab (compartido entre todas las DBs)",                                      "Sí"),
    ("GEOTAB_PASSWORD",      "contraseña Geotab",                                                                    "Sí"),
    ("ERP_SUBSCRIPTION_KEY", "clave API Gateway SAP Kaufmann (solo para build_sap_cache.py)",                        "Condicional"),
    ("AZURE_TENANT_ID",      "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx",                                                "Sí"),
    ("AZURE_CLIENT_ID",      "xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx",                                                "Sí"),
    ("RADIUS_KM",            "100  (radio de cobertura en km)",                                                      "No (default 100)"),
    ("MAX_GPS_AGE_DAYS",     "15  (max días sin ping GPS para incluir unidad)",                                      "No (default 15)"),
    ("ASSIGN_MODE",          "both  (overlap + exclusive)",                                                          "No (default both)"),
    ("LOCAL_TZ",             "America/Santiago",                                                                     "No"),
    ("TICKET_MANAGERS",      "admin@kaufmann.cl,ops@kaufmann.cl",                                                    "No"),
    ("TICKET_SLA_DAYS",      "5  (días hábiles antes de marcar ticket como vencido)",                               "No (default 5)"),
]
for i, row_data in enumerate(env_vars):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    add_table_row(t5, list(row_data), bold_first=True, bg=bg)

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 8. ESTIMACIÓN DE COSTOS MENSUALES (USD)
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Estimación de Costos Mensuales (USD)", level=1)
doc.add_paragraph(
    "Precios aproximados basados en AWS us-east-1 y Azure. "
    "No incluyen soporte premium, dominio, ni costos de transferencia de datos significativos."
)

t6 = doc.add_table(rows=1, cols=4)
t6.style = "Table Grid"
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
table_header_row(t6, ["Componente", "Especificación", "Dev/Staging (USD/mes)", "Producción (USD/mes)"])

cost_rows = [
    ("EC2 App Server",        "t3.small / t3.large",         "~$15",  "~$60"),
    ("AWS RDS PostgreSQL",    "db.t3.micro / db.t3.medium",  "~$15",  "~$55"),
    ("RDS Storage",           "20 GB / 100 GB gp3",          "~$3",   "~$12"),
    ("RDS Backup Storage",    "7 días retención",             "~$1",   "~$5"),
    ("NAT Gateway",           "Opcional en dev",              "—",     "~$32"),
    ("Elastic IP",            "IP fija para EC2",             "~$4",   "~$4"),
    ("Azure AD",              "Entra ID Free / P1",           "~$0",   "~$0–$6/usuario"),
    ("Dominio + SSL",         "Let's Encrypt (gratuito)",     "~$0",   "~$0"),
    ("",                      "TOTAL ESTIMADO",               "~$38",  "~$168+"),
]
for i, row_data in enumerate(cost_rows):
    bg = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    if row_data[0] == "":
        bg = "1F3864"
        row = t6.add_row()
        for j, v in enumerate(row_data):
            row.cells[j].text = v
            set_cell_bg(row.cells[j], bg)
            for para in row.cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    else:
        add_table_row(t6, list(row_data), bold_first=True, bg=bg, center=True)

doc.add_paragraph()
add_note(doc,
    "NAT Gateway tiene costo fijo elevado (~$32/mes). En desarrollo se puede omitir "
    "usando una IP pública directa en el EC2 con el Security Group restringido. "
    "Azure AD Free soporta hasta 50.000 usuarios/mes para autenticación básica — "
    "suficiente para este caso de uso.")
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 9. CHECKLIST DE IMPLEMENTACIÓN
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Checklist de Implementación", level=1)

phases = [
    ("Fase 1 — Infraestructura AWS", [
        "Crear VPC con subnets pública y privada",
        "Lanzar instancia EC2 t3.large (Ubuntu 22.04 LTS)",
        "Crear RDS PostgreSQL 15 en subnet privada con cifrado activado",
        "Configurar Security Groups (EC2: 443/80/22 — RDS: 5432 solo desde EC2)",
        "Habilitar extensión PostGIS en RDS:  CREATE EXTENSION IF NOT EXISTS postgis;",
        "Asignar Elastic IP a la instancia EC2",
    ]),
    ("Fase 2 — Azure AD", [
        "Crear App Registration en Azure Portal (Entra ID)",
        "Configurar Redirect URI con la URL del servidor",
        "Copiar Tenant ID y Client ID",
        "Asignar usuarios/grupos al enterprise application",
    ]),
    ("Fase 3 — Despliegue de la aplicación", [
        "Clonar repo desde github.com/ConnectDev-ai/geoworkshop-kf",
        "Instalar Python 3.11, nginx, supervisor, certbot",
        "pip install -r Scripts/requirements.txt",
        "Crear .env con todas las variables de entorno",
        "Correr pipeline inicial:  python Scripts/app.py",
        "Configurar gunicorn + supervisor para web_app.py",
        "Configurar nginx con HTTPS (certbot --nginx)",
        "Agregar cron job para pipeline diario a las 14:00 UTC",
    ]),
    ("Fase 4 — Adaptación del código para Azure AD", [
        "Reemplazar _verify_supabase_token por validación RS256 con JWKS de Azure",
        "Reemplazar Supabase JS en connect_talleres.html por MSAL.js",
        "Adaptar función authFetch() para usar tokens de MSAL",
        "Actualizar variables de entorno AZURE_TENANT_ID y AZURE_CLIENT_ID",
        "Probar login end-to-end y validación de token en el backend",
    ]),
    ("Fase 5 — Validación", [
        "Verificar que el pipeline escribe datos en RDS correctamente",
        "Verificar que /api/data responde con datos actuales",
        "Probar flujo de login con Azure AD desde el browser",
        "Probar todas las rutas del SPA (mapa, ejecutivo, detalle, tendencia, estado-flota)",
        "Verificar que el cron job del pipeline se ejecuta diariamente",
        "Revisar logs de nginx y gunicorn",
    ]),
]

for phase_name, tasks in phases:
    heading(doc, phase_name, level=2)
    for task in tasks:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run("☐  " + task)
        run.font.size = Pt(9)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 10. CONSIDERACIONES DE SEGURIDAD
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "10. Consideraciones de Seguridad", level=1)

sec_items = [
    "Nunca exponer el puerto 5432 (RDS) a internet — solo accesible desde el Security Group del EC2.",
    "El archivo .env no debe subirse al repositorio — está en .gitignore. Usar AWS Secrets Manager o Parameter Store en producción.",
    "SSL/TLS obligatorio en toda comunicación: sslmode=require en DATABASE_URL y HTTPS en nginx.",
    "El token de Azure AD se valida en cada request contra el JWKS público de Microsoft (RS256) — no se puede falsificar.",
    "El SAP ERP subscription key (ERP_SUBSCRIPTION_KEY) solo es necesario en el servidor que corre build_sap_cache.py, no en el servidor web.",
    "Restringir SSH (puerto 22) solo a IPs conocidas del equipo de desarrollo.",
    "Activar AWS CloudTrail y VPC Flow Logs para auditoría.",
    "Habilitar alertas de CloudWatch para CPU > 80% en EC2 y RDS.",
]
for item in sec_items:
    p = doc.add_paragraph(item, style="List Bullet")
    p.runs[0].font.size = Pt(9)

doc.add_paragraph()

# ── Footer ──────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    f"Documento generado automáticamente · Geoworkshop · Grupo Kaufmann · {datetime.date.today().strftime('%d/%m/%Y')}"
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

# ── Guardar ──────────────────────────────────────────────────────────────────
out_path = "Geoworkshop_Infraestructura.docx"
doc.save(out_path)
print(f"Documento generado: {out_path}")
